import sys
from pathlib import Path

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

path = Path(r"E:\素材管理软件开发文件\素材管理软件开发文档-第三阶段AI功能进行中版.docx")
doc = Document(path)
texts = [p.text for p in doc.paragraphs if p.text.strip()]

print("exists", path.exists())
print("size", path.stat().st_size)
print("paragraphs", len(doc.paragraphs))
print("tables", len(doc.tables))
print("last_text:")
for text in texts[-15:]:
    print(text[:120])
