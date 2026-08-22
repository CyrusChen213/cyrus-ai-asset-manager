from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DOC = Path(r"E:\素材管理软件开发文件\素材管理软件开发文档-第一阶段收尾版.docx")
OUT_DOC = Path(r"E:\素材管理软件开发文件\素材管理软件开发文档-第二阶段浏览器扩展进行中版.docx")
DESKTOP_DOC = Path(r"C:\Users\anzai\Desktop\素材管理软件开发文档-第二阶段浏览器扩展进行中版.docx")

FONT = "Microsoft YaHei"
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)
ACCENT = RGBColor(69, 80, 98)


def set_run(run, size=10.5, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph(paragraph, before=0, after=4, line=1.08):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, before=12 if level == 1 else 8, after=6)
    r = p.add_run(text)
    set_run(r, size=15 if level == 1 else 12, bold=True, color=ACCENT if level == 1 else TEXT)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    set_paragraph(p)
    r = p.add_run(text)
    set_run(r, size=10.5, bold=bold, color=TEXT)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    set_paragraph(p, after=3)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    r = p.add_run(f"• {text}")
    set_run(r, size=10.2, color=TEXT)
    return p


def add_check_item(doc, text, status="待验收"):
    p = doc.add_paragraph()
    set_paragraph(p, after=3)
    p.paragraph_format.left_indent = Inches(0.18)
    r1 = p.add_run(f"[{status}] ")
    set_run(r1, size=10.2, bold=True, color=ACCENT)
    r2 = p.add_run(text)
    set_run(r2, size=10.2, color=TEXT)
    return p


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["模块", "当前状态", "说明"]):
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        r = p.add_run(text)
        set_run(r, size=10, bold=True, color=TEXT)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph(p, after=0)
            r = p.add_run(text)
            set_run(r, size=9.5, color=TEXT)
    doc.add_paragraph()
    return table


def main():
    doc = Document(str(BASE_DOC))

    section = doc.sections[-1]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)

    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, before=0, after=4)
    run = title.add_run("第二阶段浏览器扩展进行中记录")
    set_run(run, size=18, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(subtitle, after=12)
    run = subtitle.add_run("更新日期：2026-06-28｜状态：功能主线已成型，进入细节验收与稳定性补强")
    set_run(run, size=10, color=MUTED)

    add_heading(doc, "二十二、第二阶段当前进度")
    add_body(
        doc,
        "第二阶段的目标是让浏览器扩展可以在常见网页中扫描当前可见素材，并把图片、动图和视频发送到桌面素材库。当前已经不只是规划状态，核心链路已经可以测试和使用。",
    )

    add_table(
        doc,
        [
            ("浮窗入口", "已完成，待验收", "扩展以网页浮窗方式工作，可拖动、关闭，并尽量在切换网页和刷新后保持可用。"),
            ("浮窗视觉", "已完成，待验收", "已改为浅色专业素材库风格，和桌面端白天模式保持一致。"),
            ("浮窗尺寸", "已完成，待验收", "右下角保留不可见热区，可拖动调整大小；位置和尺寸会记忆。"),
            ("网页扫描", "持续增强", "以当前屏幕可见素材为主，结合 DOM、懒加载字段、背景图、视频请求记录等机制。"),
            ("素材展示", "已完成，待验收", "两列瀑布流显示，尽量按真实比例展示素材，不显示多余网址和名称信息。"),
            ("收集操作", "已完成，待验收", "支持单选、多选、全选、复制选中链接、发送到桌面素材库。"),
            ("发送状态", "已完成，待验收", "素材卡片支持发送中、成功、失败状态；失败项提供重试入口。"),
            ("桌面进度", "已完成，待验收", "桌面端顶部显示批量导入进度、剩余数量、失败数量等信息。"),
            ("视频识别", "持续增强", "已支持 mp4、webm、m3u8 等常见视频线索，复杂网站仍需要继续补强。"),
        ],
    )

    add_heading(doc, "二十三、已经完成的扩展功能")
    completed = [
        "浏览器扩展浮窗可以在网页中打开，不再依赖传统小弹窗。",
        "浮窗支持拖动，并记住上次位置。",
        "浮窗支持调整大小，并记住上次尺寸；右下角不显示突兀的斜线。",
        "浮窗改为浅色模式，按钮、卡片、状态标记与桌面端风格统一。",
        "扫描结果默认不全选，避免误收集。",
        "扫描结果采用两列瀑布流，图片、动图、视频尽量按真实比例显示。",
        "素材卡片只保留必要信息：预览图、格式标记、选择状态、发送状态。",
        "支持单选、多选、全选当前扫描结果。",
        "支持复制选中素材链接。",
        "支持选择桌面端文件夹；用户上次选择的文件夹会记住。",
        "支持发送选中素材到桌面软件，成功多少就先保存多少，不要求全部成功。",
        "支持每个素材单独显示成功、失败、发送中状态。",
        "失败素材可以重试。",
        "桌面端会显示批量导入状态，动态提示总数、已完成、剩余、失败。",
        "插件刷新网页后会尝试自动重新弹出浮窗，减少用户重复打开插件的动作。",
        "扫描视频时会尽量让视频在浮窗中可播放，便于判断是否真的识别到视频。",
    ]
    for item in completed:
        add_bullet(doc, item)

    add_heading(doc, "二十四、当前已知边界和需要继续观察的问题")
    limits = [
        "不同网站的视频实现差异很大，有的网站只暴露封面图，有的网站使用分片流、临时 blob、接口延迟加载，视频识别仍需要持续补强。",
        "部分网站需要先滚动、等待懒加载完成，插件才能扫描到更多素材。",
        "网页刷新后插件会尝试重新打开，但受 Chrome 扩展权限、页面加载速度和网页自身脚本影响，仍可能出现短暂延迟。",
        "复杂瀑布流网站中，当前屏幕素材的判断需要继续调优，避免漏扫或扫到过多非当前区域内容。",
        "有声视频如果网站把音频和视频分离加载，当前阶段不强制合并音轨，先保证视频素材可收集。",
        "受版权、登录权限、平台规则保护的内容，不做破解式采集。",
    ]
    for item in limits:
        add_bullet(doc, item)

    add_heading(doc, "二十五、第二阶段验收清单")
    checks = [
        "打开桌面素材库软件，确认本地服务正常，素材库路径已经设置。",
        "在 Chrome 扩展管理页重新加载扩展，确保使用最新代码。",
        "打开一个普通图片网站，测试扫描当前屏幕素材。",
        "打开 Pinterest、花瓣等瀑布流网站，测试扫描当前屏幕素材是否足够完整。",
        "测试单选、多选、全选，再发送到桌面素材库。",
        "测试批量发送中部分成功、部分失败时，成功项是否先保存，失败项是否显示失败和重试。",
        "测试 GIF、animated WebP、普通 WebP、JPG、PNG 是否能正确识别和导入。",
        "测试 mp4、webm、m3u8 视频是否能被识别、预览和导入。",
        "测试切换保存文件夹后，再次打开插件是否记住上次文件夹。",
        "测试拖动浮窗位置、调整浮窗大小、刷新网页后位置和大小是否保留。",
        "测试关闭浮窗后，之前打开的页面是否不会残留无法关闭的旧浮窗。",
        "测试桌面端导入完成后，缩略图、视频封面、文件夹归属和来源信息是否正确。",
    ]
    for item in checks:
        add_check_item(doc, item)

    add_heading(doc, "二十六、第二阶段建议补做项")
    todo = [
        "增加更明确的“扫描范围”提示，例如当前屏幕、当前页面、滚动后再次扫描。",
        "对扫描结果增加轻量去重说明，避免同一素材因不同尺寸链接重复出现。",
        "继续增强花瓣网、Pinterest、RunningHub 等复杂网站的视频和动图识别。",
        "给插件增加更清楚的错误原因：只找到封面、视频过大、网站拒绝读取、临时链接失效等。",
        "补充来源信息在桌面端详情页的展示和筛选能力。",
        "整理插件和桌面端的中文文案，避免出现调试味较重的提示。",
        "为第二阶段做一次完整真实路径巡检：扫描、筛选、发送、失败重试、桌面查看、删除。",
    ]
    for item in todo:
        add_bullet(doc, item)

    add_heading(doc, "二十七、下一阶段建议")
    add_body(
        doc,
        "第二阶段稳定后，建议进入第三阶段：AI 自动整理。第三阶段先不要一次做太重，优先接入可配置 API，让图片自动生成内容标签、类型标签、风格标签和颜色描述，再做批量分析队列与失败重试。",
    )
    next_steps = [
        "先做 AI 配置入口：API Key、模型选择、连接测试。",
        "先做单张图片自动标签，再扩展到批量分析。",
        "标签生成结果必须允许用户编辑、删除和批量修改。",
        "AI 分析任务进入后台队列，不阻塞导入和浏览素材。",
        "保留“未分析”“整理失败”“未打标签”等入口，方便用户补处理。",
    ]
    for item in next_steps:
        add_bullet(doc, item)

    add_heading(doc, "二十八、当前建议优先级")
    priorities = [
        ("P0", "第二阶段验收巡检", "先按真实用户路径完整测试插件，不急着加入 AI，避免旧问题藏起来。"),
        ("P1", "复杂网站扫描稳定性", "继续减少漏扫、误识别、重复素材和视频下载失败。"),
        ("P1", "来源信息展示", "桌面端要更清楚地显示网页标题、来源链接、原始素材 URL。"),
        ("P2", "AI 自动标签", "第二阶段稳定后开始接入，先单张，后批量。"),
        ("P2", "Prompt 反推", "在 AI 配置和任务队列稳定后再做，支持中英对照和详细程度选择。"),
    ]
    add_table(doc, priorities)

    note = doc.add_paragraph()
    set_paragraph(note, before=4, after=0)
    r = note.add_run("备注：本文档保留第一阶段记录，并追加第二阶段真实开发进度。后续每完成一个阶段，建议继续追加“阶段验收记录”，避免开发内容和测试重点脱节。")
    set_run(r, size=10, color=MUTED)

    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOC))
    doc.save(str(DESKTOP_DOC))
    print(OUT_DOC)
    print(DESKTOP_DOC)


if __name__ == "__main__":
    main()
