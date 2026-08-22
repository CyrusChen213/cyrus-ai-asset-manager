from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


INPUT = "asset-dev-doc-base.docx"
OUTPUT = "asset-dev-doc-updated.docx"


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.first_line_indent = Pt(-6)
    run = paragraph.add_run("• " + text)
    set_run_font(run, size=10.5)


doc = Document(INPUT)

if doc.paragraphs and doc.paragraphs[-1].text.strip():
    doc.add_paragraph()

heading = doc.add_paragraph()
heading_run = heading.add_run("二十一、第一阶段收尾验收记录（2026-06-23）")
set_run_font(heading_run, size=15, bold=True, color=(31, 41, 55))

intro = doc.add_paragraph()
intro_run = intro.add_run("当前第一阶段以“本地素材库可用、导入预览稳定、基础整理闭环”为目标，已完成一次收尾补强。")
set_run_font(intro_run, size=10.5)

sections = [
    (
        "已完成补强",
        [
            "导入失败、拖拽导入失败、保存编辑副本失败、删除失败、取消重复导入失败时，均增加中文提示。",
            "整理任务失败时会即时提示，并保留“整理失败”状态，后续可以重试或忽略。",
            "提示词反推区域改为正式占位文案，明确第一阶段只记录待生成需求，第二阶段接入 AI 服务后再生成。",
            "标记待生成提示词后会给出成功反馈，避免用户误以为没有操作成功。",
            "错误通知增加低饱和红色提示样式，和当前界面风格保持一致。",
            "项目已通过 npm run build 构建验证。",
        ],
    ),
    (
        "第一阶段当前完成范围",
        [
            "本地素材库路径设置、数据库、备份目录和原始素材本地保存。",
            "图片、动图、视频导入；重复导入弹窗支持“导入 / 取消”。",
            "图片清晰缩略图、视频封面帧提取、瀑布流展示、缩略图大小调整。",
            "搜索、格式筛选、颜色筛选、标签筛选、未生成提示词、未打标签等基础查找能力。",
            "单选、多选、Shift 连选、框选、Ctrl + A、批量删除、右键删除和 Delete 删除。",
            "文件夹新建、双击重命名、删除文件夹时处理内部素材记录。",
            "双击预览、滚轮缩放、拖动、复原、自由裁剪、比例裁剪、旋转、保存为新副本。",
            "Enter 确认、Esc 取消/关闭等常用键盘交互。",
            "深色 / 浅色主题切换，中文为主的界面文案。",
        ],
    ),
    (
        "建议用户验收路径",
        [
            "新建或选择一个素材库目录，确认软件能正常进入主界面。",
            "导入 3 张不同比例图片、1 个 GIF 或 WebP 动图、1 个视频，观察缩略图和视频封面是否正常。",
            "重复导入同一张图片，确认弹窗只有“导入”和“取消”，并分别测试保留副本和撤销副本。",
            "双击图片进入预览，测试滚轮缩放、拖动、自由裁剪、比例裁剪、旋转和保存副本。",
            "使用搜索框、颜色筛选、格式筛选、标签筛选，确认筛选条件可以叠加和清除。",
            "测试单选删除、多选删除、右键删除、Delete 删除，并确认删除前有风险提示。",
            "新建文件夹、双击重命名文件夹、移动素材到文件夹，再测试删除文件夹处理方式。",
            "关闭软件后重新打开，确认素材库路径、素材列表、标签、文件夹和缩略图仍然存在。",
        ],
    ),
    (
        "暂不进入第一阶段的内容",
        [
            "浏览器扩展网页素材扫描与收集。",
            "真实 AI 图片提示词反推、批量反推和 API 配置。",
            "相似图搜索、智能文件夹、标签合并、重复素材批量清理。",
            "PSD、AI、PDF、RAW、HEIC 等更复杂格式的专业解析。",
            "安装包打包、自动更新、正式发布渠道。",
        ],
    ),
    (
        "第二阶段建议入口",
        [
            "优先做浏览器扩展：扫描当前网页素材，支持单选、多选、全选，并尽量收集最高质量原图。",
            "扩展收集时默认保存页面标题、来源链接、原始素材 URL 和收藏时间。",
            "主程序增加网页来源信息展示和来源筛选。",
            "再接入 AI 反推提示词 API，支持单张、批量、中英双语和详细程度选择。",
        ],
    ),
]

for title, items in sections:
    paragraph = doc.add_paragraph()
    title_run = paragraph.add_run(title)
    set_run_font(title_run, size=12, bold=True)
    for item in items:
        add_bullet(doc, item)

note = doc.add_paragraph()
note_run = note.add_run("备注：第一阶段可以继续小修体验问题，但不建议在验收前混入第二阶段的大功能，避免范围变大导致测试困难。")
set_run_font(note_run, size=10, italic=True)

doc.save(OUTPUT)
