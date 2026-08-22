from pathlib import Path
from shutil import copyfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

base = Path(r"E:\素材管理软件开发文件\素材管理软件开发文档-第二阶段浏览器扩展进行中版.docx")
out = Path(r"E:\素材管理软件开发文件\素材管理软件开发文档-第三阶段AI功能进行中版.docx")
desktop = Path(r"C:\Users\anzai\Desktop\素材管理软件开发文档-第三阶段AI功能进行中版.docx")

copyfile(base, out)
doc = Document(out)

# Basic style polish for appended section only through direct formatting helpers.
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(9)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color='CBD5E1'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '6')
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.bold = True
    run.font.color.rgb = RGBColor(25, 41, 70)
    if level == 1:
        run.font.size = Pt(18)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    return p


def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = 'Microsoft YaHei'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r1.font.size = Pt(10)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = 'Microsoft YaHei'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r.font.size = Pt(10)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.first_line_indent = Pt(-10)
        r = p.add_run(f'• {item}')
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r.font.size = Pt(9.5)


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color='FFFFFF')
        set_cell_shading(hdr[i], '334155')
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in('w:tcMar')
            if tc_mar is None:
                tc_mar = OxmlElement('w:tcMar')
                tc_pr.append(tc_mar)
            for side in ('top', 'left', 'bottom', 'right'):
                node = tc_mar.find(qn(f'w:{side}'))
                if node is None:
                    node = OxmlElement(f'w:{side}')
                    tc_mar.append(node)
                node.set(qn('w:w'), '90')
                node.set(qn('w:type'), 'dxa')
    doc.add_paragraph()
    return table

# Append clean page break and new section.
doc.add_page_break()
add_heading('二十八、第三阶段 AI 功能进行中记录（2026-07-01）', 1)
add_para('本次更新记录软件从“浏览器扩展收集稳定化”进入“AI 辅助整理”后的实际开发进度。当前重点已经从单纯导入、预览、删除、筛选，推进到 API 配置、AI 标签、提示词反推和批量处理。')

add_heading('1. 当前已经完成的能力', 2)
add_bullets([
    'AI 设置页：支持配置 API 地址、Key、模型名称，支持保存多个 API 配置并快速切换，可填写备注。',
    '模型识别与选择：支持拉取模型列表，提供搜索和下拉选择；对不适合识图的模型会提示用户更换。',
    '第三方/聚合 API：按 OpenAI 兼容接口设计，同时增加 RunningHub 专门配置入口，方便后续扩展更多服务。',
    'AI 标签：支持单张素材生成标签；标签可双击编辑，悬停显示删除按钮；颜色标签支持多色展示。',
    '批量 AI 标签：多选素材后可批量生成 AI 标签，逐张执行，显示进度、成功、失败，并支持重试失败项。',
    '提示词反推：支持单张图片反推中英提示词，详细程度可选；反推中有转圈动效，可取消，失败或超时会提示并可重试。',
    '批量提示词反推：多选素材后可选择简洁、中等、详细、超详细，逐张反推；视频自动跳过，成功结果即时保存，失败项可重试。',
    '颜色取样体验：素材颜色支持悬停显示十六进制色号，点击即可复制，并显示复制成功提示。',
    '多选工具条：新增导出、全选、AI 标签、批量反推等批量动作，并做了紧凑化布局。',
])

add_heading('2. 桌面端和浏览器扩展当前状态', 2)
add_table(
    ['模块', '当前状态', '备注'],
    [
        ['桌面素材库', '可日常测试', '导入、预览、瀑布流、筛选、删除、文件夹、拖拽、复制粘贴等基础链路已形成闭环。'],
        ['图片编辑', '基础可用', '支持裁剪、比例裁剪、旋转、保存副本；仍建议后续继续做边界场景测试。'],
        ['浏览器扩展', '可继续巡检', '支持浮窗、扫描当前网页素材、两列瀑布流、选择发送、成功失败状态；复杂网站仍需持续增强扫描稳定性。'],
        ['视频收集', '阶段性可用', '常见视频可以收集；对于音视频分离的网站，当前不强制合并音轨。'],
        ['AI 功能', '进行中', '已接入基础 API 配置、AI 标签、提示词反推和批量处理；后续重点是结果管理和稳定性。'],
    ],
)

add_heading('3. 仍未完成或建议补做', 2)
add_table(
    ['优先级', '事项', '说明'],
    [
        ['P0', 'AI 结果管理', '补齐复制中文 prompt、英文 prompt、重新反推、清空提示词、批量复制等常用动作。'],
        ['P0', 'AI 失败诊断', '把 401、模型不支持识图、超时、文件过大等错误改成更像产品提示的中文说明。'],
        ['P1', '插件多网站巡检', '继续测试花瓣、Pinterest、RunningHub、百度好看、普通图库和电商站，记录扫描不到或误判格式的情况。'],
        ['P1', '本地数据安全', '增加素材库数据库自动备份、恢复入口、整库导出，降低误删或数据库损坏风险。'],
        ['P1', '大量素材性能测试', '用 1000 张、5000 张素材测试瀑布流、搜索、筛选、缩略图缓存和批量操作。'],
        ['P2', '正式安装包', '制作 Windows 安装包、桌面快捷方式、插件安装说明和升级策略。'],
        ['P2', '更完整图片编辑', '后续可加入标注、简单调色、尺寸调整、格式转换等轻编辑功能。'],
    ],
)

add_heading('4. 下一步建议顺序', 2)
add_para('建议先不要马上扩很多新功能，而是把 AI 这条链路打磨顺，再做安全和打包。')
add_bullets([
    '第一步：完善 AI 结果管理，包括复制、重新生成、清空、批量复制和更清楚的失败提示。',
    '第二步：做一次第三阶段 AI 功能验收巡检，按真实操作路径测试单张、批量、失败重试、模型切换。',
    '第三步：继续插件巡检，重点测试多格式混合页面、动图、视频、受保护图片和懒加载页面。',
    '第四步：补本地数据安全能力，包括自动备份、恢复、整库导出。',
    '第五步：开始准备 Windows 打包安装和插件安装说明。',
])

add_heading('5. 当前验收清单', 2)
add_table(
    ['验收项', '检查方式', '结果记录'],
    [
        ['批量反推提示词', '多选 2-5 张图片，分别测试中等和详细，观察进度、成功、失败和重试。', '当前用户反馈暂无明显问题。'],
        ['批量 AI 标签', '多选图片执行 AI 标签，确认标签写入、失败可重试、视频自动跳过。', '需要继续抽测。'],
        ['API 配置切换', '保存多个 API 配置，切换后测试连接和反推。', '需要继续抽测。'],
        ['插件收集', '在常见网站扫描并发送图片、动图、视频到桌面端。', '阶段性可用，复杂网站继续巡检。'],
        ['桌面素材管理', '导入、筛选、删除、复制、拖拽、预览、返回位置。', '当前无高优先级问题。'],
    ],
)

add_heading('6. 备注', 2)
add_para('当前项目已经从第一阶段的“本地素材库可用”、第二阶段的“网页素材收集”，推进到第三阶段的“AI 辅助整理”。后续开发应优先保证稳定性和可解释性，让中文用户不需要理解 API 和模型细节，也能清楚知道为什么成功、为什么失败、下一步该怎么处理。')

# Save both locations.
doc.save(out)
copyfile(out, desktop)
print(out)
print(desktop)
